#!/usr/bin/env python3
from __future__ import annotations

import json
import re
import shutil
import subprocess
import sys
import textwrap
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageColor, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent.parent
PAPER = ROOT / "paper"
ARXIV_DIR = PAPER / "arxiv"
TMLR_DIR = PAPER / "tmlr"
SHARED_DIR = PAPER / "shared"
GENERATED_DIR = SHARED_DIR / "generated"
DELIVERABLES = ROOT / "deliverables"
WORD_OUTPUT = DELIVERABLES / "arxiv_preprint.docx"
WORD_OUTPUT_MIRROR = DELIVERABLES / "arxiv_preprint 9-2-2.docx"
TMLR_WORD_OUTPUT = DELIVERABLES / "tmlr_submission.docx"
WORD_TEMPLATE = SHARED_DIR / "reference.docx"
WORD_FIGURES = GENERATED_DIR / "word_assets"
ARXIV_PACKAGE = DELIVERABLES / "arxiv_submission"
TMLR_PACKAGE = DELIVERABLES / "tmlr_submission"
INTEGRITY_REPORT = DELIVERABLES / "submission_integrity_report.json"

PAPER_TITLE = "Regime Boundaries for Delayed Utility in Multi-Timescale Memory Routing"
PAPER_SUBTITLE = ""
WORD_TITLE = PAPER_TITLE
ARXIV_AUTHOR_LINES = [
    "Chandana Charitha Peddinti, University of Maryland College Park",
    "Hema Raju Barri, Johns Hopkins University",
]


SECTION_FILES = [
    ARXIV_DIR / "sections" / "abstract.tex",
    ARXIV_DIR / "sections" / "introduction.tex",
    SHARED_DIR / "sections" / "problem.tex",
    SHARED_DIR / "sections" / "system.tex",
    ARXIV_DIR / "sections" / "benchmarks.tex",
    ARXIV_DIR / "sections" / "popqa.tex",
    ARXIV_DIR / "sections" / "public_freshqa.tex",
    ARXIV_DIR / "sections" / "mquake.tex",
    ARXIV_DIR / "sections" / "uniedit.tex",
    ARXIV_DIR / "sections" / "limitations.tex",
    ARXIV_DIR / "sections" / "reproducibility.tex",
    ARXIV_DIR / "sections" / "appendix.tex",
]

TMLR_SECTION_FILES = [
    TMLR_DIR / "sections" / "abstract.tex",
    TMLR_DIR / "sections" / "introduction.tex",
    SHARED_DIR / "sections" / "problem.tex",
    SHARED_DIR / "sections" / "system.tex",
    TMLR_DIR / "sections" / "benchmarks.tex",
    TMLR_DIR / "sections" / "popqa.tex",
    TMLR_DIR / "sections" / "public_freshqa.tex",
    TMLR_DIR / "sections" / "mquake.tex",
    TMLR_DIR / "sections" / "uniedit.tex",
    TMLR_DIR / "sections" / "limitations.tex",
    TMLR_DIR / "sections" / "reproducibility.tex",
    TMLR_DIR / "sections" / "appendix.tex",
]

SVG_TO_PNG = {
    GENERATED_DIR / "popqa_retrieval_calls_reduction.svg": WORD_FIGURES / "popqa_retrieval_calls_reduction.png",
    GENERATED_DIR / "popqa_action_distribution.svg": WORD_FIGURES / "popqa_action_distribution.png",
    GENERATED_DIR / "freshness_stale_answer_rate.svg": WORD_FIGURES / "freshness_stale_answer_rate.png",
    GENERATED_DIR / "freshness_action_distribution.svg": WORD_FIGURES / "freshness_action_distribution.png",
    GENERATED_DIR / "fig_system_overview.svg": WORD_FIGURES / "fig_system_overview.png",
    GENERATED_DIR / "fig_public_main_calibration_curve.svg": WORD_FIGURES / "fig_public_main_calibration_curve.png",
    GENERATED_DIR / "fig_public_audit_verdicts.svg": WORD_FIGURES / "fig_public_audit_verdicts.png",
}


@dataclass
class BibEntry:
    key: str
    fields: dict[str, str]


class CitationManager:
    def __init__(self, entries: dict[str, BibEntry]) -> None:
        self.entries = entries
        self.order: list[str] = []
        self.index: dict[str, int] = {}

    def cite(self, keys: list[str]) -> str:
        numbers: list[str] = []
        for key in keys:
            key = key.strip()
            if not key:
                continue
            if key not in self.index:
                self.order.append(key)
                self.index[key] = len(self.order)
            numbers.append(str(self.index[key]))
        return "[" + ",".join(numbers) + "]"

    def ordered_entries(self) -> list[tuple[int, BibEntry]]:
        return [(self.index[key], self.entries[key]) for key in self.order if key in self.entries]


def parse_bibtex(path: Path) -> dict[str, BibEntry]:
    text = path.read_text()
    entries: dict[str, BibEntry] = {}
    for match in re.finditer(r"@(\w+)\{([^,]+),", text):
        key = match.group(2).strip()
        start = match.end()
        depth = 1
        end = start
        while end < len(text):
            ch = text[end]
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    break
            end += 1
        body = text[start:end]
        fields = {
            field.group(1).strip(): field.group(2).strip().strip("{}")
            for field in re.finditer(r"(\w+)\s*=\s*\{((?:[^{}]|\{[^{}]*\})*)\}\s*,?", body, re.S)
        }
        entries[key] = BibEntry(key=key, fields=fields)
    return entries


def ensure_word_template(path: Path) -> None:
    if path.exists():
        return
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for style_name, size in [("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
    doc.save(path)


def font_for_size(size: int) -> ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/SFNS.ttf",
    ]
    for candidate in candidates:
        candidate_path = Path(candidate)
        if candidate_path.exists():
            return ImageFont.truetype(str(candidate_path), size)
    return ImageFont.load_default()


def parse_svg_length(value: str | None, full: int) -> int:
    if value is None:
        return 0
    value = value.strip()
    if value.endswith("%"):
        return int(float(value[:-1]) * full / 100.0)
    return int(float(value))


def render_simple_svg(svg_path: Path, png_path: Path) -> None:
    png_path.parent.mkdir(parents=True, exist_ok=True)
    root = ET.fromstring(svg_path.read_text())
    width = parse_svg_length(root.attrib.get("width", "960"), 960)
    height = parse_svg_length(root.attrib.get("height", "480"), 480)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    ns = "{http://www.w3.org/2000/svg}"

    for element in root:
        tag = element.tag.replace(ns, "")
        if tag == "rect":
            x = parse_svg_length(element.attrib.get("x", "0"), width)
            y = parse_svg_length(element.attrib.get("y", "0"), height)
            w = parse_svg_length(element.attrib.get("width", str(width)), width)
            h = parse_svg_length(element.attrib.get("height", str(height)), height)
            fill = element.attrib.get("fill")
            if fill:
                rx = int(float(element.attrib.get("rx", "0")))
                if rx > 0 and hasattr(draw, "rounded_rectangle"):
                    draw.rounded_rectangle([x, y, x + w, y + h], radius=rx, fill=ImageColor.getrgb(fill))
                else:
                    draw.rectangle([x, y, x + w, y + h], fill=ImageColor.getrgb(fill))
        elif tag == "line":
            points = [
                parse_svg_length(element.attrib.get("x1", "0"), width),
                parse_svg_length(element.attrib.get("y1", "0"), height),
                parse_svg_length(element.attrib.get("x2", "0"), width),
                parse_svg_length(element.attrib.get("y2", "0"), height),
            ]
            stroke = ImageColor.getrgb(element.attrib.get("stroke", "#000"))
            draw.line(points, fill=stroke, width=2)
        elif tag == "text":
            x = parse_svg_length(element.attrib.get("x", "0"), width)
            y = parse_svg_length(element.attrib.get("y", "0"), height)
            fill = ImageColor.getrgb(element.attrib.get("fill", "#000"))
            font_size = int(float(element.attrib.get("font-size", "12")))
            font = font_for_size(font_size)
            text = "".join(element.itertext())
            anchor = element.attrib.get("text-anchor", "start")
            bbox = draw.textbbox((0, 0), text, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            if anchor == "middle":
                x -= text_width // 2
            elif anchor == "end":
                x -= text_width
            draw.text((x, y - text_height), text, fill=fill, font=font)

    image.save(png_path)


def generate_png_assets() -> None:
    for svg_path, png_path in SVG_TO_PNG.items():
        render_simple_svg(svg_path, png_path)


def load_text(path: Path) -> str:
    return path.read_text()


def resolve_input_path(current_dir: Path, raw_rel: str) -> Path:
    candidates = [
        (current_dir / raw_rel).resolve(),
        (ARXIV_DIR / raw_rel).resolve(),
        (PAPER / raw_rel).resolve(),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
        if not candidate.suffix:
            tex_candidate = candidate.with_suffix(".tex")
            if tex_candidate.exists():
                return tex_candidate
    candidate = candidates[0]
    return candidate


def strip_latex_markup(text: str, citations: CitationManager) -> str:
    def replace_cite(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        return citations.cite(keys)

    def unwrap(cmd: str, value: str) -> str:
        return value

    text = re.sub(r"Table~\\ref\{[^}]+\}", "the table below", text)
    text = re.sub(r"Figure~\\ref\{[^}]+\}", "the figure below", text)
    text = re.sub(r"Algorithm~\\ref\{[^}]+\}", "the routing policy below", text)
    text = re.sub(r"\\cite\w*\{([^}]+)\}", replace_cite, text)
    for pattern in [r"\\texttt\{([^{}]*)\}", r"\\emph\{([^{}]*)\}", r"\\textbf\{([^{}]*)\}", r"\\text\{([^{}]*)\}", r"\\mathrm\{([^{}]*)\}", r"\\textrm\{([^{}]*)\}"]:
        while re.search(pattern, text):
            text = re.sub(pattern, lambda m: unwrap("", m.group(1)), text)
    text = re.sub(r"\\\((.*?)\\\)", lambda m: m.group(1), text)
    text = re.sub(r"\\\[(.*?)\\\]", lambda m: m.group(1), text)
    text = text.replace(r"\begin{quote}", "").replace(r"\end{quote}", "")
    text = text.replace(r"\noindent", "")
    text = text.replace(r"\_", "_").replace(r"\%", "%")
    text = text.replace("~", " ")
    text = text.replace(r"\today", "")
    text = text.replace(r"\textbackslash", "\\")
    text = text.replace(r"\mathcal{A}", "A")
    text = text.replace(r"\in", "in")
    text = text.replace(r"\rightarrow", "->")
    text = text.replace(r"\times", "x")
    text = text.replace(r"\beta", "β")
    text = text.replace(r"\{", "{").replace(r"\}", "}")
    text = re.sub(r"\^\{([^}]+)\}", r"^\1", text)
    text = re.sub(r"_\{([^}]+)\}", r"_\1", text)
    text = text.replace("β V_{future}", "β V_future")
    text = text.replace("β V_{\\text{future}}", "β V_future")
    text = text.replace("$V_{future}$", "V_future")
    text = text.replace("$V_{\\text{future}}$", "V_future")
    text = text.replace("V_{future}", "V_future")
    text = text.replace(r"V_{\mathrm{future}}", "V_future")
    text = text.replace(r"V_{\text{future}}", "V_future")
    text = text.replace("$", "")
    text = text.replace("`", "")
    text = re.sub(r"\\label\{[^}]+\}", "", text)
    text = re.sub(r"\\ref\{[^}]+\}", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_table_rows(path: Path, citations: CitationManager) -> list[list[str]]:
    rows: list[list[str]] = []
    for raw_line in path.read_text().splitlines():
        line = raw_line.strip()
        if not line or line.startswith(r"\toprule") or line.startswith(r"\midrule") or line.startswith(r"\bottomrule"):
            continue
        if not line.endswith(r"\\"):
            continue
        line = line[:-2].strip()
        if "&" in line:
            cells = [strip_latex_markup(cell.strip(), citations) for cell in line.split("&")]
        else:
            cells = [strip_latex_markup(line, citations)]
        rows.append(cells)
    return rows


def parse_blocks(text: str, current_dir: Path) -> list[dict]:
    blocks: list[dict] = []
    lines = text.splitlines()
    i = 0
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            paragraph = " ".join(line.strip() for line in paragraph_lines if line.strip())
            if paragraph:
                blocks.append({"type": "paragraph", "text": paragraph})
            paragraph_lines = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            i += 1
            continue
        if stripped in {r"\clearpage", r"\newpage"}:
            flush_paragraph()
            blocks.append({"type": "page_break"})
            i += 1
            continue
        if stripped.startswith(r"\section{"):
            flush_paragraph()
            title = stripped[len(r"\section{") : -1]
            blocks.append({"type": "heading", "level": 1, "text": title})
            i += 1
            continue
        if stripped.startswith(r"\subsection{"):
            flush_paragraph()
            title = stripped[len(r"\subsection{") : -1]
            blocks.append({"type": "heading", "level": 2, "text": title})
            i += 1
            continue
        if stripped.startswith(r"\paragraph{"):
            flush_paragraph()
            title_end = stripped.index("}")
            title = stripped[len(r"\paragraph{") : title_end]
            remainder = stripped[title_end + 1 :].strip()
            blocks.append({"type": "paragraph_heading", "title": title, "text": remainder})
            i += 1
            continue
        if stripped.startswith(r"\begin{itemize}"):
            flush_paragraph()
            items: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(r"\end{itemize}"):
                item_line = lines[i].strip()
                if item_line.startswith(r"\item"):
                    items.append(item_line[len(r"\item") :].strip())
                i += 1
            blocks.append({"type": "bullet_list", "items": items})
            i += 1
            continue
        if stripped.startswith(r"\begin{center}"):
            flush_paragraph()
            center_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(r"\end{center}"):
                center_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "centered", "text": " ".join(center_lines)})
            i += 1
            continue
        if stripped.startswith(r"\["):
            flush_paragraph()
            equation_lines = [stripped[2:].strip()]
            i += 1
            while i < len(lines) and r"\]" not in lines[i]:
                equation_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                equation_lines.append(lines[i].split(r"\]")[0].strip())
            equation = " ".join(part for part in equation_lines if part)
            blocks.append({"type": "equation", "text": equation})
            i += 1
            continue
        if stripped.startswith(r"\begin{table}"):
            flush_paragraph()
            caption = ""
            table_file: Path | None = None
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(r"\end{table}"):
                inner = lines[i].strip()
                if inner.startswith(r"\input{"):
                    rel = inner[len(r"\input{") : -1]
                    table_file = resolve_input_path(current_dir, rel)
                elif inner.startswith(r"\caption{"):
                    caption = inner[len(r"\caption{") : -1]
                i += 1
            blocks.append({"type": "table", "caption": caption, "path": table_file})
            i += 1
            continue
        if stripped.startswith(r"\begin{figure}"):
            flush_paragraph()
            caption = ""
            image_file: Path | None = None
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(r"\end{figure}"):
                inner = lines[i].strip()
                if inner.startswith(r"\includegraphics"):
                    match = re.search(r"\{([^}]+)\}", inner)
                    if match:
                        image_file = resolve_input_path(current_dir, match.group(1))
                elif inner.startswith(r"\caption{"):
                    caption = inner[len(r"\caption{") : -1]
                i += 1
            blocks.append({"type": "figure", "caption": caption, "path": image_file})
            i += 1
            continue
        if stripped.startswith(r"\input{"):
            flush_paragraph()
            rel = stripped[len(r"\input{") : -1]
            include_path = resolve_input_path(current_dir, rel)
            include_text = load_text(include_path)
            blocks.extend(parse_blocks(include_text, include_path.parent))
            i += 1
            continue
        if stripped.startswith(r"\appendix") or stripped.startswith(r"\bibliographystyle") or stripped.startswith(r"\bibliography") or stripped.startswith(r"\begin{abstract}") or stripped.startswith(r"\end{abstract}"):
            flush_paragraph()
            i += 1
            continue
        paragraph_lines.append(stripped)
        i += 1
    flush_paragraph()
    return blocks


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(text)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def apply_doc_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(doc: Document, rows: list[list[str]], caption: str, citations: CitationManager) -> None:
    if not rows:
        return
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = True
    cap.add_run(strip_caption(caption, citations)).italic = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for ridx, row in enumerate(rows):
        for cidx, cell_text in enumerate(row):
            set_cell_text(table.cell(ridx, cidx), cell_text)
        set_row_cant_split(table.rows[ridx])
        if ridx == 0:
            set_repeat_table_header(table.rows[0])
    doc.add_paragraph()


def strip_caption(caption: str, citations: CitationManager) -> str:
    return strip_latex_markup(caption, citations)


def add_figure(doc: Document, image_path: Path, caption: str, citations: CitationManager) -> None:
    if image_path is None or not image_path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(strip_caption(caption, citations)).italic = True
    doc.add_paragraph()


def build_references(doc: Document, citations: CitationManager) -> None:
    doc.add_heading("References", level=1)
    for number, entry in citations.ordered_entries():
        fields = entry.fields
        authors = fields.get("author", "").replace(" and ", ", ")
        title = fields.get("title", "")
        venue = fields.get("booktitle") or fields.get("journal", "")
        year = fields.get("year", "")
        doi = fields.get("doi", "")
        url = fields.get("url", "")
        link = ""
        if doi:
            link = doi if doi.startswith("http") else f"https://doi.org/{doi}"
        elif url:
            link = url
        parts = [f"[{number}] {authors}.", title + ".", venue + ".", year + "."]
        if link:
            parts.append(link)
        doc.add_paragraph(" ".join(part for part in parts if part and part != "."))


def build_word_document(
    *,
    output_path: Path,
    author_lines: list[str],
    section_files: list[Path],
    abstract_path: Path,
) -> None:
    ensure_word_template(WORD_TEMPLATE)
    entries = parse_bibtex(SHARED_DIR / "references.bib")
    citations = CitationManager(entries)
    doc = Document(str(WORD_TEMPLATE))
    apply_doc_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(WORD_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(18)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run("\n".join(author_lines))

    doc.add_paragraph()
    doc.add_heading("Abstract", level=1)
    abstract_text = strip_latex_markup(load_text(abstract_path), citations)
    doc.add_paragraph(abstract_text)

    for section_path in section_files[1:]:
        blocks = parse_blocks(load_text(section_path), section_path.parent)
        for block in blocks:
            if block["type"] == "heading":
                heading = doc.add_heading(strip_latex_markup(block["text"], citations), level=block["level"])
                heading.paragraph_format.keep_with_next = True
            elif block["type"] == "paragraph":
                doc.add_paragraph(strip_latex_markup(block["text"], citations))
            elif block["type"] == "paragraph_heading":
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.keep_with_next = True
                title_run = paragraph.add_run(strip_latex_markup(block["title"], citations) + " ")
                title_run.bold = True
                if block["text"]:
                    paragraph.add_run(strip_latex_markup(block["text"], citations))
            elif block["type"] == "bullet_list":
                for item in block["items"]:
                    doc.add_paragraph(strip_latex_markup(item, citations), style="List Bullet")
            elif block["type"] == "centered":
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(strip_latex_markup(block["text"], citations))
            elif block["type"] == "equation":
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                eq = strip_latex_markup(block["text"], citations)
                paragraph.add_run(eq).italic = True
            elif block["type"] == "page_break":
                add_page_break(doc)
            elif block["type"] == "table" and block["path"]:
                rows = parse_table_rows(block["path"], citations)
                add_table(doc, rows, block["caption"], citations)
            elif block["type"] == "figure" and block["path"]:
                add_figure(doc, block["path"], block["caption"], citations)

    build_references(doc, citations)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def package_paper_submission(source_dir: Path, target_dir: Path, paper_dir_name: str) -> None:
    if target_dir.exists():
        shutil.rmtree(target_dir)
    (target_dir / paper_dir_name / "sections").mkdir(parents=True, exist_ok=True)
    (target_dir / "shared" / "sections").mkdir(parents=True, exist_ok=True)
    (target_dir / "shared" / "generated").mkdir(parents=True, exist_ok=True)

    shutil.copy2(source_dir / "main.tex", target_dir / paper_dir_name / "main.tex")
    for path in (source_dir / "sections").glob("*.tex"):
        shutil.copy2(path, target_dir / paper_dir_name / "sections" / path.name)
    for path in (SHARED_DIR / "sections").glob("*.tex"):
        shutil.copy2(path, target_dir / "shared" / "sections" / path.name)
    shutil.copytree(GENERATED_DIR, target_dir / "shared" / "generated", dirs_exist_ok=True)
    shutil.copy2(SHARED_DIR / "preamble.tex", target_dir / "shared" / "preamble.tex")
    shutil.copy2(SHARED_DIR / "references.bib", target_dir / "shared" / "references.bib")
    shutil.copy2(SHARED_DIR / "artifact_availability.tex", target_dir / "shared" / "artifact_availability.tex")
    shutil.copy2(SHARED_DIR / "claims_to_artifacts.md", target_dir / "shared" / "claims_to_artifacts.md")
    shutil.copy2(ROOT / "paper_run_manifest_index.json", target_dir / "paper_run_manifest_index.json")
    shutil.copy2(ROOT / "DATASET_MANIFEST.md", target_dir / "DATASET_MANIFEST.md")


def package_arxiv_submission() -> None:
    package_paper_submission(ARXIV_DIR, ARXIV_PACKAGE, "arxiv")

    readme = textwrap.dedent(
        """\
        # arXiv Submission Package

        This package contains the source for the full arXiv preprint.

        ## Included

        - `arxiv/main.tex`
        - `arxiv/sections/*.tex`
        - `shared/preamble.tex`
        - `shared/references.bib`
        - `shared/sections/*.tex`
        - `shared/generated/*` for frozen tables and figure assets
        - `paper_run_manifest_index.json`
        - `DATASET_MANIFEST.md`

        ## Notes

        - This package is source-complete relative to the named arXiv manuscript.
        - Only bundles listed in `paper_run_manifest_index.json` are submission-facing and reportable.
        - The package exposes four evidence surfaces inside the manuscript and package notes: the main evidence surface, the appendix-only supplementary diagnostic, the manual audit bundle, and benchmark reliability notes.
        - The main evidence surface is the registry-listed set of roots under `artifacts/tmlr_official/`; inspect `manifest.json`, `aggregate_table.md`, `run_rows.jsonl`, `significance.json`, and `bootstrap.json` in the cited bundle first.
        - The appendix-only supplementary diagnostic is `artifacts/freshness_v1/`.
        - The manual audit bundle is `artifacts/tmlr_official/freshqa_leakage_audit_manual_v1/`.
        - The benchmark reliability notes are `artifacts/tmlr_official/mquake_reliability_note_v1/` and `artifacts/tmlr_official/uniedit_reliability_note_v1/`.
        - Use `paper/shared/claims_to_artifacts.md`, `paper_run_manifest_index.json`, and `DATASET_MANIFEST.md` together to verify any headline claim.
        - A public archive may be posted alongside or after preprint release, but this package does not depend on any external repository link.
        - The environment that built this package did not have `pdflatex` or `latexmk`, so PDF compilation was not performed here.
        - The manuscript is backed by the frozen official roots under `artifacts/tmlr_official/` plus the appendix-only supplementary artifact `artifacts/freshness_v1/`.
        """
    )
    checklist = textwrap.dedent(
        """\
        # Publication Checklist

        - [x] Headline claims point only to frozen official roots under `artifacts/tmlr_official/`.
        - [x] Bundled controlled-update context appears only as appendix-only arXiv supplementary context.
        - [x] Citations come from `shared/references.bib`.
        - [x] Figures and tables are copied from frozen generated assets.
        - [x] The public benchmark caveat remains visible in the manuscript.
        - [x] No visible LaTeX source artifacts should remain in the Word export.
        - [ ] Compile-check `arxiv/main.tex` in a TeX environment with `pdflatex` or `latexmk`.
        - [ ] Validate final PDF fonts and machine readability before submission.
        """
    )
    (ARXIV_PACKAGE / "README.md").write_text(readme)
    (ARXIV_PACKAGE / "PUBLICATION_CHECKLIST.md").write_text(checklist)


def package_tmlr_submission() -> None:
    package_paper_submission(TMLR_DIR, TMLR_PACKAGE, "tmlr")
    (TMLR_PACKAGE / "arxiv" / "sections").mkdir(parents=True, exist_ok=True)
    for path in (ARXIV_DIR / "sections").glob("*.tex"):
        if path.name == "reproducibility.tex":
            continue
        shutil.copy2(path, TMLR_PACKAGE / "arxiv" / "sections" / path.name)
    readme = textwrap.dedent(
        """\
        # TMLR Submission Package

        This package contains the anonymous TMLR-facing source manuscript and supplementary records for the frozen evidence bundles.

        ## Included

        - `tmlr/main.tex`
        - `arxiv/sections/*.tex` as the shared anonymous paper body
        - `shared/preamble.tex`
        - `shared/references.bib`
        - `shared/sections/*.tex`
        - `shared/generated/*` for paper tables and figure assets
        - `paper_run_manifest_index.json`
        - `DATASET_MANIFEST.md`

        ## Notes

        - This package is anonymous by construction and should remain unlinked to any author-identified archive in the submitted PDF.
        - Only bundles listed in `paper_run_manifest_index.json` are submission-facing and reportable.
        - The environment that built this package did not have `pdflatex` or `latexmk`, so PDF compilation was not performed here.
        - Supplementary evidence should remain anonymous and point only to the frozen roots listed in `shared/claims_to_artifacts.md`.
        """
    )
    checklist = textwrap.dedent(
        """\
        # Anonymity Checklist

        - [x] No author names or affiliations in `tmlr/main.tex`.
        - [x] No acknowledgements or funding text in the anonymous manuscript.
        - [x] No direct links to author-identified archives in the anonymous package.
        - [x] Citations come from `shared/references.bib`.
        - [x] Figures and tables are copied from frozen generated assets.
        - [ ] Compile-check `tmlr/main.tex` in a TeX environment with `pdflatex` or `latexmk`.
        - [ ] Review supplementary filenames and external links before OpenReview upload.
        """
    )
    manifest = textwrap.dedent(
        """\
        # Supplement Manifest

        Anonymous supplementary material should include:

        - frozen traces, manifests, tables, and audits from the evidence roots listed in `shared/claims_to_artifacts.md`
        - benchmark construction notes from `docs/freshness_snapshot_protocol.md` and `docs/journal_reproducibility.md`
        - the public-track provenance bundle and manual audit bundle
        - exact rerun commands for all tables reported in the manuscript
        """
    )
    (TMLR_PACKAGE / "README.md").write_text(readme)
    (TMLR_PACKAGE / "ANONYMITY_CHECKLIST.md").write_text(checklist)
    (TMLR_PACKAGE / "SUPPLEMENT_MANIFEST.md").write_text(manifest)


def write_delivery_docs() -> None:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    delivery_note = textwrap.dedent(
        """\
        # Delivery Note

        - Editable full manuscript: `arxiv_preprint.docx`
        - Legacy mirrored snapshot: `arxiv_preprint 9-2-2.docx`
        - Editable anonymous TMLR manuscript: `tmlr_submission.docx`
        - Direct arXiv source package: `arxiv_submission/`
        - Anonymous TMLR source package: `tmlr_submission/`

        The Word documents mirror the named arXiv preprint and the anonymous TMLR manuscript. The LaTeX packages are the authoritative direct-submission paths for arXiv and TMLR.
        """
    )
    (DELIVERABLES / "README.md").write_text(delivery_note)


def extract_word_text(path: Path) -> str:
    if not path.exists():
        raise RuntimeError(f"expected generated Word output at {path}, but it does not exist")
    completed = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )
    if completed.returncode != 0:
        raise RuntimeError(
            "failed to extract generated Word text with textutil.\n"
            f"stdout:\n{completed.stdout}\n"
            f"stderr:\n{completed.stderr}"
        )
    return completed.stdout


def validate_arxiv_word_output(path: Path) -> None:
    text = extract_word_text(path)
    banned_patterns = {
        "future-value-collapse": r"\bV_future" + r"rewards\b",
        "leading-fficial-caption": r"(?<![A-Za-z])f" + r"ficial\b",
        "algorithm-input-row": r"\bIn" + r"puts:\b",
        "inline-latex-residue": r"\\\(",
    }
    remaining = [label for label, pattern in banned_patterns.items() if re.search(pattern, text)]
    if remaining:
        raise RuntimeError(
            "generated arXiv Word export contains banned fragments: "
            + ", ".join(remaining)
        )
    required_fragments = [
        "score_t(a) =",
        "Main evidence surface",
        "Appendix-only supplementary diagnostic",
        "Manual audit",
        "Reliability notes",
        "How to verify",
    ]
    missing = [fragment for fragment in required_fragments if fragment not in text]
    if missing:
        raise RuntimeError(
            "generated arXiv Word export is missing required verification text: "
            + ", ".join(missing)
        )


def run_integrity_audit(*, include_packages: bool) -> None:
    targets = [
        "paper/arxiv",
        "paper/tmlr",
        "paper/shared/claims_to_artifacts.md",
        "paper/shared/generated/README.md",
        "paper/shared/artifact_availability.tex",
        "docs/journal_reproducibility.md",
        "deliverables/PUBLICATION_STATUS.md",
        "deliverables/README.md",
        "paper_run_manifest_index.json",
    ]
    if include_packages:
        targets.extend(
            [
                "deliverables/arxiv_submission",
                "deliverables/tmlr_submission",
            ]
        )
    command = [
        sys.executable,
        str(ROOT / "scripts" / "audit_submission_integrity.py"),
        "--report",
        str(INTEGRITY_REPORT),
    ]
    for target in targets:
        command.extend(["--path", target])
    completed = subprocess.run(command, cwd=ROOT, capture_output=True, text=True, check=False)
    if completed.returncode != 0:
        raise RuntimeError(
            "submission-integrity audit failed.\n"
            f"stdout:\n{completed.stdout}\n"
            f"stderr:\n{completed.stderr}"
        )


def main() -> None:
    run_integrity_audit(include_packages=False)
    generate_png_assets()
    build_word_document(
        output_path=WORD_OUTPUT,
        author_lines=ARXIV_AUTHOR_LINES,
        section_files=SECTION_FILES,
        abstract_path=ARXIV_DIR / "sections" / "abstract.tex",
    )
    build_word_document(
        output_path=TMLR_WORD_OUTPUT,
        author_lines=["Anonymous Submission"],
        section_files=TMLR_SECTION_FILES,
        abstract_path=TMLR_DIR / "sections" / "abstract.tex",
    )
    validate_arxiv_word_output(WORD_OUTPUT)
    shutil.copy2(WORD_OUTPUT, WORD_OUTPUT_MIRROR)
    package_arxiv_submission()
    package_tmlr_submission()
    write_delivery_docs()
    run_integrity_audit(include_packages=True)
    summary = {
        "word_output": str(WORD_OUTPUT.relative_to(ROOT)),
        "word_output_mirror": str(WORD_OUTPUT_MIRROR.relative_to(ROOT)),
        "tmlr_word_output": str(TMLR_WORD_OUTPUT.relative_to(ROOT)),
        "submission_package": str(ARXIV_PACKAGE.relative_to(ROOT)),
        "tmlr_package": str(TMLR_PACKAGE.relative_to(ROOT)),
        "template": str(WORD_TEMPLATE.relative_to(ROOT)),
        "png_assets": [str(path.relative_to(ROOT)) for path in WORD_FIGURES.glob("*.png")],
        "integrity_report": str(INTEGRITY_REPORT.relative_to(ROOT)),
    }
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
